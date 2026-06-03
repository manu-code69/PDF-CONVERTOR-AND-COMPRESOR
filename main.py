import os, io, shutil, uuid, gc
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

TARGET_SIZES = {
    "50": 50 * 1024,
    "100": 100 * 1024,
    "500": 500 * 1024,
    "0": None
}

def compress_pdf_to_target(input_path: str, output_path: str, target_bytes):
    import fitz
    doc = fitz.open(input_path)

    if target_bytes is None:
        doc.save(output_path, garbage=4, deflate=True, clean=True)
        doc.close()
        return

    for dpi in [150, 96, 72, 48, 32]:
        tmp = output_path + ".tmp.pdf"
        new_doc = fitz.open()
        for page in doc:
            mat = fitz.Matrix(dpi / 72, dpi / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_bytes = pix.tobytes("jpeg")
            img_rect = fitz.Rect(0, 0, pix.width, pix.height)
            page_pdf = fitz.open()
            new_page = page_pdf.new_page(width=pix.width, height=pix.height)
            new_page.insert_image(img_rect, stream=img_bytes)
            new_doc.insert_pdf(page_pdf)
            page_pdf.close()
        new_doc.save(tmp, garbage=4, deflate=True, clean=True)
        new_doc.close()
        size = os.path.getsize(tmp)
        if size <= target_bytes or dpi == 32:
            shutil.move(tmp, output_path)
            break
        os.remove(tmp)
    doc.close()


def image_to_pdf(input_path: str, output_path: str, target_bytes):
    from PIL import Image
    import fitz

    img = Image.open(input_path).convert("RGB")

    quality = 85
    if target_bytes:
        for q in [85, 70, 55, 40, 25, 10]:
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=q, optimize=True)
            if buf.tell() < target_bytes - 1024 or q == 10:
                quality = q
                break

    tmp_jpg = input_path + "_tmp.jpg"
    img.save(tmp_jpg, format="JPEG", quality=quality, optimize=True)

    img_doc = fitz.open(tmp_jpg)
    pdfbytes = img_doc.convert_to_pdf()
    img_doc.close()

    with open(output_path, "wb") as f:
        f.write(pdfbytes)

    try:
        os.remove(tmp_jpg)
    except Exception:
        pass

    if target_bytes and os.path.getsize(output_path) > target_bytes:
        tmp = output_path + ".recompress.pdf"
        compress_pdf_to_target(output_path, tmp, target_bytes)
        shutil.move(tmp, output_path)


def docx_to_pdf(input_path: str, output_path: str, target_bytes):
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    doc = Document(input_path)
    styles = getSampleStyleSheet()

    custom_normal = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontSize=11, leading=16, spaceAfter=6)
    custom_h1 = ParagraphStyle('CustomH1', parent=styles['Heading1'], fontSize=16, leading=20, spaceAfter=10)
    custom_h2 = ParagraphStyle('CustomH2', parent=styles['Heading2'], fontSize=13, leading=18, spaceAfter=8)

    story = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue
        style_name = para.style.name.lower()
        if 'heading 1' in style_name or 'title' in style_name:
            story.append(Paragraph(text, custom_h1))
        elif 'heading 2' in style_name:
            story.append(Paragraph(text, custom_h2))
        else:
            story.append(Paragraph(text, custom_normal))

    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        if data:
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('PADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 12))

    pdf_doc = SimpleDocTemplate(output_path, pagesize=A4,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch)
    pdf_doc.build(story)

    if target_bytes and os.path.getsize(output_path) > target_bytes:
        tmp = output_path + ".recompress.pdf"
        compress_pdf_to_target(output_path, tmp, target_bytes)
        shutil.move(tmp, output_path)


@app.get("/", response_class=HTMLResponse)
async def root():
    with open("templates/index.html", encoding="utf-8") as f:
        return f.read()


@app.post("/convert")
async def convert(file: UploadFile = File(...), target_size: str = Form("0")):
    uid = uuid.uuid4().hex[:8]
    ext = Path(file.filename).suffix.lower()
    allowed = {".png", ".jpg", ".jpeg", ".pdf", ".docx"}

    if ext not in allowed:
        raise HTTPException(400, f"Unsupported file type: {ext}")

    input_path = str(UPLOAD_DIR / f"{uid}_input{ext}")
    output_name = Path(file.filename).stem + "_converted.pdf"
    output_path = str(OUTPUT_DIR / f"{uid}_output.pdf")

    with open(input_path, "wb") as f:
        f.write(await file.read())

    target_bytes = TARGET_SIZES.get(target_size)
    original_size = os.path.getsize(input_path)

    try:
        if ext in {".png", ".jpg", ".jpeg"}:
            image_to_pdf(input_path, output_path, target_bytes)
        elif ext == ".docx":
            docx_to_pdf(input_path, output_path, target_bytes)
        elif ext == ".pdf":
            compress_pdf_to_target(input_path, output_path, target_bytes)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Conversion failed: {str(e)}")
    finally:
        gc.collect()
        try:
            if os.path.exists(input_path):
                os.remove(input_path)
        except Exception:
            pass

    if not os.path.exists(output_path):
        raise HTTPException(500, "Output file not generated")

    final_size = os.path.getsize(output_path)

    return {
        "download_url": f"/download/{uid}",
        "filename": output_name,
        "original_size": original_size,
        "final_size": final_size,
    }


@app.get("/download/{uid}")
async def download(uid: str):
    path = OUTPUT_DIR / f"{uid}_output.pdf"
    if not path.exists():
        raise HTTPException(404, "File not found")
    return FileResponse(str(path), media_type="application/pdf",
                        headers={"Content-Disposition": 'attachment; filename="converted.pdf"'})